"""
Report generation — HTML, PDF, DOCX, and JSON formats.
"""

import datetime
import json
from pathlib import Path


# ─────────────────────────────────────────────────────────────────────────────
# Threat level scoring
# ─────────────────────────────────────────────────────────────────────────────

_LEVEL_MAP = {
    "CRITICAL": ("CRITICAL", "#dc2626", (220, 38, 38)),
    "HIGH":     ("HIGH",     "#ea580c", (234, 88, 12)),
    "MEDIUM":   ("MEDIUM",   "#ca8a04", (202, 138, 4)),
    "LOW":      ("LOW",      "#16a34a", (22, 163, 74)),
}


def _threat_level(analysis: dict) -> tuple:
    """
    Derive threat level from multiple sources in priority order:
      1. Structured AI output (most reliable — explicit classification).
      2. Threat intel (known malware = at least HIGH).
      3. Static indicators (scoring heuristics as fallback).
    """
    ai      = analysis.get("ai_analysis", {})
    static  = analysis.get("static", {})
    intel   = analysis.get("intel", {})

    # ── 1. Structured AI output — authoritative if available ─────────────────
    structured = ai.get("structured") or {}
    ai_level   = (structured.get("threat_level") or "").upper().strip()
    if ai_level in _LEVEL_MAP:
        return _LEVEL_MAP[ai_level]

    # ── 2. Known malware from threat intel ───────────────────────────────────
    if intel.get("_summary", {}).get("known_malware"):
        return _LEVEL_MAP["HIGH"]

    # ── 3. Static heuristic scoring (fallback) ───────────────────────────────
    packer  = static.get("packer", {})
    pe      = static.get("pe", {})
    entropy = static.get("entropy", {})
    script  = static.get("script", {})
    pdf     = static.get("pdf", {})
    office  = static.get("office", {}) or static.get("ole", {})

    score = 0

    # PE indicators
    if packer and packer.get("suspicious"):
        score += 2
    if pe:
        susp = len(pe.get("suspicious_imports", []))
        score += min(susp // 2, 3)
        if pe.get("has_tls"):
            score += 1
    if entropy and entropy.get("suspicious"):
        score += 2

    # Script indicators
    if script and not script.get("error"):
        obs = script.get("obfuscation_score", 0)
        score += min(obs // 2, 3)
        sp = script.get("suspicious_patterns", {})
        critical_ps = {"iex", "amsi_bypass", "reflection", "encoded_command",
                       "eval", "activex", "wscript_shell", "exec_eval", "base64_decode"}
        hit_critical = len(critical_ps & set(sp.keys()))
        score += min(hit_critical, 3)
        if script.get("decoded_payloads"):
            score += 2

    # PDF indicators
    if pdf:
        summary = pdf.get("summary", {}) or pdf.get("raw_structure", {})
        if summary.get("has_javascript"):
            score += 2
        if summary.get("has_embedded_files") or summary.get("has_auto_actions"):
            score += 1
        if summary.get("suspicious"):
            score += 1

    # Office indicators
    if office:
        ole_s = office.get("summary", {})
        if ole_s.get("has_macros"):
            score += 1
        if ole_s.get("suspicious_macros"):
            score += 2
        if ole_s.get("cve_indicators"):
            score += 3

    # XOR deobfuscation results — decoded payloads means obfuscated content
    xor = static.get("xor_deobfuscation", {})
    if xor.get("found_payloads"):
        score += 2

    # IOC signals (from any source)
    iocs = static.get("strings", {}).get("iocs", {})
    script_iocs = (script or {}).get("iocs", {})
    all_iocs = {**iocs, **script_iocs}
    if all_iocs.get("urls") or all_iocs.get("ips"):
        score += 1
    if all_iocs.get("api_calls") or all_iocs.get("suspicious_keywords"):
        score += 1

    if score >= 7:
        return _LEVEL_MAP["CRITICAL"]
    if score >= 4:
        return _LEVEL_MAP["HIGH"]
    if score >= 2:
        return _LEVEL_MAP["MEDIUM"]
    return _LEVEL_MAP["LOW"]


# ─────────────────────────────────────────────────────────────────────────────
# Master entry point
# ─────────────────────────────────────────────────────────────────────────────

def generate_report(analysis: dict, output_path: str, fmt: str = "html") -> str:
    """
    Generate a report in the requested format.

    fmt: "html" | "pdf" | "docx" | "json"
    Returns the path to the written file.
    """
    fmt = fmt.lower().lstrip(".")
    p = Path(output_path)

    # Normalise extension
    suffix_map = {"html": ".html", "pdf": ".pdf", "docx": ".docx", "json": ".json"}
    ext = suffix_map.get(fmt, ".html")
    if p.suffix.lower() != ext:
        p = p.with_suffix(ext)

    p.parent.mkdir(parents=True, exist_ok=True)

    if fmt == "json":
        return _write_json(analysis, str(p))
    if fmt == "pdf":
        return _write_pdf(analysis, str(p))
    if fmt == "docx":
        return _write_docx(analysis, str(p))
    # default: html
    return _write_html(analysis, str(p))


def generate_all(analysis: dict, base_path: str) -> dict:
    """
    Generate all 4 report formats from a base path (without extension).
    Returns dict of {fmt: file_path}.
    """
    base = Path(base_path)
    results = {}
    for fmt in ("html", "pdf", "docx", "json"):
        out = str(base.with_suffix(f".{fmt}"))
        try:
            results[fmt] = generate_report(analysis, out, fmt=fmt)
        except Exception as e:
            results[fmt] = f"ERROR: {e}"
    return results


# ─────────────────────────────────────────────────────────────────────────────
# JSON
# ─────────────────────────────────────────────────────────────────────────────

def _write_json(analysis: dict, path: str) -> str:
    with open(path, "w", encoding="utf-8") as f:
        json.dump(analysis, f, indent=2, ensure_ascii=False, default=str)
    return path


# ─────────────────────────────────────────────────────────────────────────────
# HTML
# ─────────────────────────────────────────────────────────────────────────────

def _write_html(analysis: dict, path: str) -> str:
    try:
        from jinja2 import Environment, FileSystemLoader, select_autoescape
        template_dir = Path(__file__).parent / "templates"
        env = Environment(
            loader=FileSystemLoader(str(template_dir)),
            autoescape=select_autoescape(["html"]),
        )
        template = env.get_template("report.html")
        threat_level, threat_color, _ = _threat_level(analysis)
        ctx = {
            "meta":         analysis.get("meta", {}),
            "file_info":    analysis.get("file_info", {}),
            "static":       analysis.get("static", {}),
            "dynamic":      analysis.get("dynamic"),
            "ai_analysis":  analysis.get("ai_analysis", {}),
            "threat_level": threat_level,
            "threat_color": threat_color,
            "generated_at": datetime.datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC"),
        }
        html = template.render(**ctx)
        Path(path).write_text(html, encoding="utf-8")
        return path
    except Exception as e:
        # Fallback to inline HTML
        return _write_html_fallback(analysis, path)


def _write_html_fallback(analysis: dict, path: str) -> str:
    """Inline HTML when Jinja2 template fails."""
    threat_level, threat_color, _ = _threat_level(analysis)
    meta      = analysis.get("meta", {})
    file_info = analysis.get("file_info", {})
    static    = analysis.get("static", {})
    ai        = analysis.get("ai_analysis", {})
    hashes    = file_info.get("hashes", {})

    def esc(s):
        return str(s).replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")

    lines = [f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><title>Malyze Report — {esc(file_info.get('name',''))}</title>
<style>
*{{box-sizing:border-box;margin:0;padding:0}}
body{{background:#0d1117;color:#c9d1d9;font-family:'Segoe UI',system-ui,sans-serif;font-size:14px;line-height:1.6;padding:24px}}
.wrap{{max-width:1100px;margin:0 auto}}
h1{{color:#58a6ff;font-size:22px;letter-spacing:2px;text-transform:uppercase}}
h2{{color:#58a6ff;font-size:15px;text-transform:uppercase;letter-spacing:1px;margin:20px 0 10px;padding-bottom:6px;border-bottom:1px solid #30363d}}
h3{{color:#8b949e;font-size:12px;text-transform:uppercase;letter-spacing:1px;margin:12px 0 6px}}
.card{{background:#161b22;border:1px solid #30363d;border-radius:8px;padding:20px;margin-bottom:20px}}
.badge{{display:inline-block;padding:2px 8px;border-radius:12px;font-size:11px;font-weight:bold}}
.threat{{font-size:22px;font-weight:bold;letter-spacing:3px;color:{threat_color}}}
table{{width:100%;border-collapse:collapse;font-size:13px}}
th,td{{padding:7px 12px;border:1px solid #30363d;text-align:left}}
th{{background:#1c2128;color:#8b949e;font-size:11px;text-transform:uppercase}}
pre{{background:#0d1117;border:1px solid #30363d;border-radius:6px;padding:14px;font-family:monospace;font-size:12px;white-space:pre-wrap;word-break:break-all}}
.red{{color:#f85149}} .orange{{color:#e3b341}} .green{{color:#3fb950}}
.mono{{font-family:monospace}}
footer{{text-align:center;color:#8b949e;font-size:12px;margin-top:32px;padding-top:16px;border-top:1px solid #30363d}}
</style></head><body><div class="wrap">

<div class="card">
<h1>&#x1F6E1; Malyze — Malware Analysis Report</h1>
<p style="margin-top:12px"><b>Analyst:</b> {esc(meta.get('analyst','N/A'))} &nbsp;|&nbsp;
<b>Org:</b> {esc(meta.get('org','N/A'))} &nbsp;|&nbsp; <b>Date:</b> {esc(meta.get('timestamp','N/A'))}</p>
<div style="margin-top:16px" class="threat">THREAT: {threat_level}</div>
</div>

<div class="card">
<h2>&#x1F4C4; File Information</h2>
<table>
<tr><th>Property</th><th>Value</th></tr>
<tr><td>Name</td><td class="mono">{esc(file_info.get('name','N/A'))}</td></tr>
<tr><td>Type</td><td>{esc(file_info.get('type','N/A'))}</td></tr>
<tr><td>Size</td><td>{hashes.get('size',0):,} bytes</td></tr>
<tr><td>MD5</td><td class="mono">{esc(hashes.get('md5','N/A'))}</td></tr>
<tr><td>SHA1</td><td class="mono">{esc(hashes.get('sha1','N/A'))}</td></tr>
<tr><td>SHA256</td><td class="mono">{esc(hashes.get('sha256','N/A'))}</td></tr>
</table>
</div>"""]

    # AI Analysis
    ai_text = esc(ai.get("analysis","")) if not ai.get("error") else f'<span class="red">Error: {esc(ai.get("error",""))}</span>'
    lines.append(f"""<div class="card">
<h2>&#x1F916; AI Analysis ({esc(ai.get('model',''))})</h2>
<pre>{ai_text}</pre>
</div>""")

    # Entropy
    entropy = static.get("entropy", {})
    if entropy:
        e_cls = "red" if entropy.get("suspicious") else "green"
        lines.append(f"""<div class="card">
<h2>&#x1F4CA; Entropy Analysis</h2>
<table>
<tr><td>Overall Entropy</td><td class="{e_cls}"><b>{entropy.get('overall_entropy')}</b> / 8.0</td></tr>
<tr><td>Classification</td><td>{esc(entropy.get('classification',''))}</td></tr>
<tr><td>Suspicious</td><td class="{e_cls}">{entropy.get('suspicious',False)}</td></tr>
<tr><td>High Entropy Blocks</td><td>{entropy.get('high_entropy_blocks',0)} / {entropy.get('total_blocks',0)}</td></tr>
</table>
</div>""")

    # Packer
    packer = static.get("packer", {})
    if packer:
        pkrs = packer.get("detected_packers", [])
        pkr_html = " ".join(f'<span class="badge" style="background:rgba(248,81,73,.15);color:#f85149;border:1px solid rgba(248,81,73,.4)">{esc(p)}</span>' for p in pkrs) if pkrs else '<span class="badge" style="background:rgba(63,185,80,.15);color:#3fb950">None detected</span>'
        lines.append(f"""<div class="card">
<h2>&#x1F4E6; Packer Detection</h2>{pkr_html}</div>""")

    # PE
    pe = static.get("pe", {})
    if pe and not pe.get("error"):
        susp = pe.get("suspicious_imports", [])
        susp_html = "".join(f'<span class="badge" style="background:rgba(248,81,73,.1);color:#f85149;border:1px solid rgba(248,81,73,.3);margin:2px">{esc(i)}</span>' for i in susp)
        sec_rows = "".join(
            f'<tr><td class="mono">{esc(s["name"])}</td>'
            f'<td class="mono">{esc(s.get("virtual_address",""))}</td>'
            f'<td>{s.get("virtual_size",0):,}</td>'
            f'<td>{s.get("raw_size",0):,}</td>'
            f'<td class="{"red" if s.get("suspicious") else "green"}">{s.get("entropy",0):.3f}</td>'
            f'<td>{", ".join(s.get("flags",[]))}</td></tr>'
            for s in pe.get("sections", [])
        )
        imp_rows = "".join(
            f'<tr><td class="mono">{esc(dll)}</td><td>{len(funcs)}</td>'
            f'<td style="font-size:11px;font-family:monospace">{esc(", ".join(funcs[:6]))}{"…" if len(funcs)>6 else ""}</td></tr>'
            for dll, funcs in pe.get("imports", {}).items()
        )
        lines.append(f"""<div class="card">
<h2>&#x1F9F1; PE Analysis</h2>
<table>
<tr><th>Field</th><th>Value</th></tr>
<tr><td>Machine</td><td>{esc(pe.get('machine',''))}</td></tr>
<tr><td>Compile Time</td><td>{esc(pe.get('timestamp',''))}</td></tr>
<tr><td>Subsystem</td><td>{esc(pe.get('subsystem',''))}</td></tr>
<tr><td>Entry Point</td><td class="mono">{esc(pe.get('entry_point',''))}</td></tr>
<tr><td>Image Base</td><td class="mono">{esc(pe.get('image_base',''))}</td></tr>
<tr><td>Is DLL</td><td>{pe.get('is_dll',False)}</td></tr>
<tr><td>Is 64-bit</td><td>{pe.get('is_64bit',False)}</td></tr>
<tr><td>Has TLS Callbacks</td><td class="{'orange' if pe.get('has_tls') else ''}">{pe.get('has_tls',False)}</td></tr>
<tr><td>Digital Signature</td><td>{pe.get('has_signature',False)}</td></tr>
</table>
<h3>Sections</h3>
<table>
<tr><th>Name</th><th>Virt. Addr</th><th>V.Size</th><th>Raw Size</th><th>Entropy</th><th>Flags</th></tr>
{sec_rows}
</table>
{"<h3 class='red'>Suspicious Imports (" + str(len(susp)) + ")</h3><div style='margin-top:6px'>" + susp_html + "</div>" if susp else ""}
{"<h3>Import Table</h3><table><tr><th>DLL</th><th>Count</th><th>Functions</th></tr>" + imp_rows + "</table>" if imp_rows else ""}
</div>""")

    # Script analysis
    script = static.get("script", {})
    if script and not script.get("error"):
        obs = script.get("obfuscation_score", 0)
        obs_color = "red" if obs >= 6 else "orange" if obs >= 3 else "green"
        sp_html = "".join(
            f'<span class="badge" style="background:rgba(248,81,73,.15);color:#f85149;border:1px solid rgba(248,81,73,.4);margin:2px">{esc(k)}</span>'
            for k in script.get("suspicious_patterns", {})
        )
        dp_html = ""
        for dp in script.get("decoded_payloads", [])[:5]:
            dp_html += f'<h3>{esc(dp.get("type","decoded"))}</h3><pre>{esc(str(dp.get("decoded",""))[:1000])}</pre>'
        script_ioc_html = ""
        for cat, items in script.get("iocs", {}).items():
            col = "#f85149" if cat in ("urls","ips","suspicious_keywords") else "#e3b341"
            pills = " ".join(f'<span class="badge" style="background:#1c2128;border:1px solid #30363d;color:{col};margin:2px;font-family:monospace">{esc(str(i))}</span>' for i in items[:20])
            script_ioc_html += f"<h3>{esc(cat.replace('_',' ').upper())}</h3><div style='margin-bottom:8px'>{pills}</div>"
        lines.append(f"""<div class="card">
<h2>&#x1F4DC; Script Analysis — {esc(script.get('file_type',''))} (Static / No Execution)</h2>
<table>
<tr><td>Line Count</td><td>{script.get('line_count',0):,}</td></tr>
<tr><td>Obfuscation Score</td><td class="{obs_color}"><b>{obs}/15</b></td></tr>
<tr><td>Indicators</td><td>{esc(', '.join(script.get('obfuscation_indicators',[])) or 'None')}</td></tr>
</table>
{"<h3>Suspicious Patterns</h3><div style='margin-top:4px'>" + sp_html + "</div>" if sp_html else ""}
{script_ioc_html}
{("<h3>Decoded Payloads</h3>" + dp_html) if dp_html else ""}
<h3>Content Sample</h3>
<pre>{esc(chr(10).join(script.get('strings_sample',[])[:80]))}</pre>
</div>""")

    # PDF analysis
    pdf = static.get("pdf", {})
    if pdf:
        summary = pdf.get("summary", {}) or pdf.get("raw_structure", {})
        if summary:
            susp_kw = summary.get("suspicious_keywords", [])
            urls_f  = summary.get("urls", [])
            susp_kw_pills = " ".join(f'<span class="badge" style="background:rgba(248,81,73,.15);color:#f85149;border:1px solid rgba(248,81,73,.4)">{esc(k)}</span>' for k in susp_kw)
            url_pills = " ".join(f'<span class="badge" style="background:#1c2128;border:1px solid #30363d;color:#f85149;font-family:monospace">{esc(u)}</span>' for u in urls_f[:20])
            lines.append(f"""<div class="card">
<h2>&#x1F4C3; PDF Analysis</h2>
<table>
<tr><td>Pages</td><td>{esc(str(summary.get('page_count','N/A')))}</td></tr>
<tr><td>JavaScript</td><td class="{'red' if summary.get('has_javascript') else ''}">{summary.get('has_javascript',False)}</td></tr>
<tr><td>Embedded Files</td><td class="{'orange' if summary.get('has_embedded_files') else ''}">{summary.get('has_embedded_files',False)}</td></tr>
<tr><td>Auto Actions (/OpenAction /AA)</td><td class="{'orange' if summary.get('has_auto_actions') else ''}">{summary.get('has_auto_actions',False)}</td></tr>
<tr><td>Suspicious</td><td class="{'red' if summary.get('suspicious') else 'green'}">{summary.get('suspicious',False)}</td></tr>
</table>
{"<h3 class='red'>Suspicious Keywords</h3><div style='margin-top:4px'>" + susp_kw_pills + "</div>" if susp_kw_pills else ""}
{"<h3>URLs Found</h3><div style='margin-top:4px'>" + url_pills + "</div>" if url_pills else ""}
</div>""")

    # Office analysis
    office = static.get("office", {}) or static.get("ole", {})
    if office:
        ole_s = office.get("summary", {})
        macros = (office.get("oletools", {}) or {}).get("macros", []) or office.get("macros", [])
        if ole_s or macros:
            macro_html = ""
            for m in macros[:5]:
                hits = m.get("suspicious_hits", [])
                hit_pills = " ".join(f'<span class="badge" style="background:rgba(248,81,73,.15);color:#f85149;border:1px solid rgba(248,81,73,.4)">{esc(h)}</span>' for h in hits)
                macro_html += f"<h3>Macro: {esc(m.get('filename',''))}</h3>{hit_pills or 'No suspicious hits'}<pre style='margin-top:8px'>{esc(m.get('code','')[:1000])}</pre>"
            lines.append(f"""<div class="card">
<h2>&#x1F4CE; Office Document Analysis</h2>
<table>
<tr><td>Has Macros</td><td class="{'red' if ole_s.get('has_macros') else 'green'}">{ole_s.get('has_macros',False)}</td></tr>
<tr><td>Suspicious Macros</td><td class="{'red' if ole_s.get('suspicious_macros') else ''}">{ole_s.get('suspicious_macros',False)}</td></tr>
<tr><td>External URLs</td><td>{len(ole_s.get('external_urls',[]))}</td></tr>
<tr><td>CVE Indicators</td><td class="{'red' if ole_s.get('cve_indicators') else ''}">{esc(', '.join(ole_s.get('cve_indicators',[])) or 'None')}</td></tr>
<tr><td>IOC Count</td><td>{ole_s.get('ioc_count',0)}</td></tr>
</table>
{macro_html}
</div>""")

    # CAPA capabilities
    capa = static.get("capa", {})
    if capa:
        caps = capa.get("capabilities", [])
        total_matched = capa.get("total_rules_matched", len(caps))
        if caps:
            rows_html = ""
            for c in caps[:60]:
                ns   = esc(c.get("namespace", ""))
                name = esc(c.get("name", ""))
                atk_list = c.get("attack", [])
                atk = ", ".join(
                    a.get("technique", str(a)) if isinstance(a, dict) else str(a)
                    for a in atk_list
                )
                rows_html += (
                    f"<tr><td style='font-family:monospace'>{name}</td>"
                    f"<td style='color:#8b949e'>{ns}</td>"
                    f"<td style='color:#f85149'>{esc(atk)}</td></tr>"
                )
            lines.append(f"""<div class="card">
<h2>&#x1F9E0; CAPA Capabilities ({total_matched} rules matched)</h2>
<table>
<tr><th>Rule</th><th>Namespace</th><th>ATT&amp;CK</th></tr>
{rows_html}
</table>
</div>""")
        elif capa.get("text"):
            lines.append(f"""<div class="card">
<h2>&#x1F9E0; CAPA Capabilities</h2>
<pre style="font-size:12px">{esc(capa['text'][:3000])}</pre>
</div>""")

    # Strings / IOCs
    strings_data = static.get("strings", {})
    iocs = strings_data.get("iocs", {})
    if iocs or strings_data.get("strings"):
        ioc_html = ""
        color_map = {"urls":"red","ips":"red","api_calls":"orange","suspicious_keywords":"red","domains":"orange"}
        for cat, items in iocs.items():
            col = color_map.get(cat, "")
            style = f'color:{"#f85149" if col=="red" else "#e3b341" if col=="orange" else "#58a6ff"}'
            pills = " ".join(f'<span class="badge" style="background:#1c2128;border:1px solid #30363d;{style};margin:2px;font-family:monospace">{esc(str(i))}</span>' for i in items[:25])
            ioc_html += f"<h3>{esc(cat.replace('_',' ').upper())}</h3><div style='margin-bottom:8px'>{pills}</div>"
        sample = "\n".join(esc(s) for s in strings_data.get("strings", [])[:200])
        lines.append(f"""<div class="card">
<h2>&#x1F50D; Strings &amp; IOCs ({strings_data.get('total',0)} total via {esc(strings_data.get('source',''))})</h2>
{ioc_html}
<h3>String Sample</h3>
<pre>{sample}</pre>
</div>""")

    # Disasm
    disasm = static.get("disassembly", {})
    if disasm and not disasm.get("error"):
        blocks_html = ""
        for bname, block in (disasm.get("blocks") or {}).items():
            if block.get("error"):
                continue
            rows = "".join(
                f'<div style="display:flex;gap:16px;font-size:12px;font-family:monospace;padding:2px 0">'
                f'<span style="color:#79c0ff;min-width:130px">{esc(i["address"])}</span>'
                f'<span style="color:#ff7b72;min-width:80px">{esc(i["mnemonic"])}</span>'
                f'<span>{esc(i["op_str"])}</span></div>'
                for i in block.get("instructions", [])[:80]
            )
            blocks_html += f'<h3>{esc(bname)} @ {esc(block.get("va",""))}</h3><div style="background:#0d1117;border:1px solid #30363d;border-radius:6px;padding:14px;overflow-x:auto">{rows}</div>'
        lines.append(f"""<div class="card">
<h2>&#x1F4BB; Disassembly ({esc(disasm.get('arch',''))})</h2>
<p style="color:#8b949e;margin-bottom:12px">Image Base: <span class="mono">{esc(disasm.get('image_base',''))}</span> &nbsp;|&nbsp; Entry RVA: <span class="mono">{esc(disasm.get('entry_point_rva',''))}</span></p>
{blocks_html}
</div>""")

    lines.append(f"""<footer>Generated by <b>Malyze</b> — AI-Powered Malware Analysis Framework<br>
Analyst: {esc(meta.get('analyst',''))} | {datetime.datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')}</footer>
</div></body></html>""")

    Path(path).write_text("\n".join(lines), encoding="utf-8")
    return path


# ─────────────────────────────────────────────────────────────────────────────
# PDF (ReportLab)
# ─────────────────────────────────────────────────────────────────────────────

def _write_pdf(analysis: dict, path: str) -> str:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
            HRFlowable, KeepTogether,
        )
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
    except ImportError:
        raise ImportError("reportlab not installed. Run: pip install reportlab")

    threat_level, _, threat_rgb = _threat_level(analysis)
    meta      = analysis.get("meta", {})
    file_info = analysis.get("file_info", {})
    static    = analysis.get("static", {})
    ai        = analysis.get("ai_analysis", {})
    hashes    = file_info.get("hashes", {})
    dynamic   = analysis.get("dynamic")

    threat_color = colors.Color(threat_rgb[0]/255, threat_rgb[1]/255, threat_rgb[2]/255)

    # Styles
    styles = getSampleStyleSheet()
    style_normal = ParagraphStyle("Normal2", parent=styles["Normal"], fontSize=9, leading=13)
    style_h1 = ParagraphStyle("H1", fontSize=18, textColor=colors.HexColor("#1a56db"),
                               spaceAfter=4, fontName="Helvetica-Bold")
    style_h2 = ParagraphStyle("H2", fontSize=12, textColor=colors.HexColor("#1a56db"),
                               spaceBefore=14, spaceAfter=4, fontName="Helvetica-Bold",
                               borderPadding=(0, 0, 3, 0))
    style_h3 = ParagraphStyle("H3", fontSize=10, textColor=colors.grey,
                               spaceBefore=8, spaceAfter=3, fontName="Helvetica-Bold")
    style_mono = ParagraphStyle("Mono", fontName="Courier", fontSize=8, leading=11,
                                 wordWrap="CJK")
    style_threat = ParagraphStyle("Threat", fontSize=20, fontName="Helvetica-Bold",
                                   textColor=threat_color, spaceAfter=8)
    style_caption = ParagraphStyle("Caption", fontSize=8, textColor=colors.grey, leading=11)

    def tbl_style(header_bg=colors.HexColor("#1a56db")):
        return TableStyle([
            ("BACKGROUND",  (0, 0), (-1, 0), header_bg),
            ("TEXTCOLOR",   (0, 0), (-1, 0), colors.white),
            ("FONTNAME",    (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE",    (0, 0), (-1, -1), 8),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8f9fa")]),
            ("GRID",        (0, 0), (-1, -1), 0.3, colors.HexColor("#dee2e6")),
            ("VALIGN",      (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING",  (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING",(0, 0), (-1, -1), 4),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING",(0, 0), (-1, -1), 6),
        ])

    doc = SimpleDocTemplate(
        path, pagesize=A4,
        leftMargin=1.8*cm, rightMargin=1.8*cm,
        topMargin=2*cm, bottomMargin=2*cm,
        title=f"Malyze Report — {file_info.get('name','')}",
    )

    story = []
    W = A4[0] - 3.6*cm  # usable width

    def hr():
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#dee2e6"), spaceAfter=6))

    # ── Cover ──
    story.append(Paragraph("MALYZE", style_h1))
    story.append(Paragraph("AI-Powered Malware Analysis Report", style_normal))
    story.append(Spacer(1, 8))
    story.append(Paragraph(f"THREAT LEVEL: {threat_level}", style_threat))
    hr()

    meta_data = [
        ["Analyst",      meta.get("analyst", "N/A")],
        ["Organization", meta.get("org", "N/A")],
        ["Date",         meta.get("timestamp", "N/A")],
        ["Tool",         meta.get("tool", "Malyze")],
        ["Generated",    datetime.datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC")],
    ]
    t = Table(meta_data, colWidths=[3*cm, W-3*cm])
    t.setStyle(tbl_style(colors.HexColor("#374151")))
    story.append(t)
    story.append(Spacer(1, 12))

    # ── File Info ──
    story.append(Paragraph("File Information", style_h2))
    fi_data = [
        ["Property", "Value"],
        ["Name",      file_info.get("name", "N/A")],
        ["Type",      file_info.get("type", "N/A")],
        ["Extension", file_info.get("extension", "N/A")],
        ["Size",      f"{hashes.get('size',0):,} bytes"],
        ["MD5",       hashes.get("md5", "N/A")],
        ["SHA1",      hashes.get("sha1", "N/A")],
        ["SHA256",    hashes.get("sha256", "N/A")],
        ["Magic Bytes", file_info.get("magic_bytes", "N/A")],
    ]
    t = Table(fi_data, colWidths=[3.5*cm, W-3.5*cm])
    t.setStyle(tbl_style())
    story.append(t)
    story.append(Spacer(1, 10))

    # ── AI Analysis ──
    story.append(Paragraph("AI-Powered Threat Analysis", style_h2))
    if ai.get("error"):
        story.append(Paragraph(f"Error: {ai['error']}", style_normal))
    elif ai.get("analysis"):
        for line in ai["analysis"].splitlines():
            line = line.strip()
            if not line:
                story.append(Spacer(1, 4))
                continue
            if line.startswith("##"):
                story.append(Paragraph(line.lstrip("#").strip(), style_h3))
            elif line.startswith("#"):
                story.append(Paragraph(line.lstrip("#").strip(), style_h2))
            else:
                story.append(Paragraph(line, style_normal))
    story.append(Spacer(1, 10))

    # ── Entropy ──
    entropy = static.get("entropy", {})
    if entropy:
        story.append(Paragraph("Entropy Analysis", style_h2))
        e_color = colors.HexColor("#dc2626") if entropy.get("suspicious") else colors.HexColor("#16a34a")
        ent_data = [
            ["Metric", "Value"],
            ["Overall Entropy", str(entropy.get("overall_entropy", "N/A"))],
            ["Classification",  entropy.get("classification", "N/A")],
            ["Suspicious",      str(entropy.get("suspicious", False))],
            ["High Entropy Blocks", f"{entropy.get('high_entropy_blocks',0)} / {entropy.get('total_blocks',0)}"],
        ]
        t = Table(ent_data, colWidths=[5*cm, W-5*cm])
        t.setStyle(tbl_style())
        story.append(t)
        story.append(Spacer(1, 10))

    # ── Packer ──
    packer = static.get("packer", {})
    if packer:
        story.append(Paragraph("Packer / Protector Detection", style_h2))
        pkrs = packer.get("detected_packers", [])
        if pkrs:
            story.append(Paragraph("Detected: " + ", ".join(pkrs), ParagraphStyle(
                "PkrAlert", parent=style_normal, textColor=colors.HexColor("#dc2626"), fontName="Helvetica-Bold"
            )))
        else:
            story.append(Paragraph("No known packers detected.", style_normal))
        story.append(Spacer(1, 10))

    # ── PE ──
    pe = static.get("pe", {})
    if pe and not pe.get("error"):
        story.append(Paragraph("PE File Analysis", style_h2))
        pe_data = [
            ["Field", "Value"],
            ["Machine",      pe.get("machine", "N/A")],
            ["Timestamp",    pe.get("timestamp", "N/A")],
            ["Subsystem",    pe.get("subsystem", "N/A")],
            ["Entry Point",  pe.get("entry_point", "N/A")],
            ["Image Base",   pe.get("image_base", "N/A")],
            ["Is DLL",       str(pe.get("is_dll", False))],
            ["Is 64-bit",    str(pe.get("is_64bit", False))],
            ["Has TLS",      str(pe.get("has_tls", False))],
            ["Has Signature",str(pe.get("has_signature", False))],
        ]
        t = Table(pe_data, colWidths=[4*cm, W-4*cm])
        t.setStyle(tbl_style())
        story.append(Spacer(1, 6))

        # Sections
        story.append(Paragraph("Sections", style_h3))
        sec_data = [["Name", "Virt. Addr", "V.Size", "Raw Size", "Entropy", "Flags"]]
        for sec in pe.get("sections", []):
            row_style_extra = [colors.HexColor("#fef2f2") if sec.get("suspicious") else colors.white]
            sec_data.append([
                sec.get("name", "?"),
                sec.get("virtual_address", ""),
                f"{sec.get('virtual_size',0):,}",
                f"{sec.get('raw_size',0):,}",
                f"{sec.get('entropy',0):.3f}",
                ", ".join(sec.get("flags", [])[:3]),
            ])
        t2 = Table(sec_data, colWidths=[2*cm, 2.5*cm, 2.2*cm, 2.2*cm, 2*cm, None])
        t2.setStyle(tbl_style())
        story.append(t2)

        susp = pe.get("suspicious_imports", [])
        if susp:
            story.append(Spacer(1, 6))
            story.append(Paragraph(f"Suspicious Imports ({len(susp)})", ParagraphStyle(
                "SuspHdr", parent=style_h3, textColor=colors.HexColor("#dc2626")
            )))
            story.append(Paragraph(", ".join(susp[:60]), ParagraphStyle(
                "SuspList", parent=style_mono, textColor=colors.HexColor("#dc2626")
            )))

        all_imports = pe.get("imports", {})
        if all_imports:
            story.append(Spacer(1, 6))
            story.append(Paragraph("Import Table", style_h3))
            imp_data = [["DLL", "Count", "Functions (sample)"]]
            for dll, funcs in list(all_imports.items())[:20]:
                imp_data.append([dll, str(len(funcs)), ", ".join(funcs[:5]) + ("…" if len(funcs)>5 else "")])
            t3 = Table(imp_data, colWidths=[4*cm, 1.5*cm, W-5.5*cm])
            t3.setStyle(tbl_style())
            story.append(t3)

        story.append(Spacer(1, 10))

    # ── IOCs ──
    strings_data = static.get("strings", {})
    iocs = strings_data.get("iocs", {})
    if iocs:
        story.append(Paragraph(f"Strings & IOCs ({strings_data.get('total',0)} total)", style_h2))
        for cat, items in iocs.items():
            story.append(Paragraph(cat.replace("_"," ").upper(), style_h3))
            story.append(Paragraph(", ".join(str(i) for i in items[:30]), style_mono))
        story.append(Spacer(1, 10))

    # ── Disassembly ──
    disasm = static.get("disassembly", {})
    if disasm and not disasm.get("error"):
        story.append(Paragraph(f"Disassembly ({disasm.get('arch','')})", style_h2))
        for bname, block in (disasm.get("blocks") or {}).items():
            if block.get("error"):
                continue
            story.append(Paragraph(f"{bname} @ {block.get('va','')}", style_h3))
            lines_asm = []
            for instr in block.get("instructions", [])[:60]:
                lines_asm.append(f"{instr['address']}  {instr['mnemonic']:<10} {instr['op_str']}")
            story.append(Paragraph("\n".join(lines_asm), style_mono))
        story.append(Spacer(1, 10))

    # ── Dynamic ──
    if dynamic:
        story.append(Paragraph("Dynamic Analysis", style_h2))
        exec_r = dynamic.get("execution", {})
        if exec_r:
            dyn_data = [
                ["Field", "Value"],
                ["PID",       str(exec_r.get("pid", "N/A"))],
                ["Exit Code", str(exec_r.get("exit_code", "N/A"))],
                ["Timed Out", str(exec_r.get("timed_out", False))],
            ]
            t = Table(dyn_data, colWidths=[4*cm, W-4*cm])
            t.setStyle(tbl_style())
            story.append(t)
        story.append(Spacer(1, 10))

    doc.build(story)
    return path


# ─────────────────────────────────────────────────────────────────────────────
# DOCX (python-docx)
# ─────────────────────────────────────────────────────────────────────────────

def _write_docx(analysis: dict, path: str) -> str:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    threat_level, _, threat_rgb = _threat_level(analysis)
    meta      = analysis.get("meta", {})
    file_info = analysis.get("file_info", {})
    static    = analysis.get("static", {})
    ai        = analysis.get("ai_analysis", {})
    hashes    = file_info.get("hashes", {})
    dynamic   = analysis.get("dynamic")

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    def set_cell_bg(cell, hex_color: str):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:fill"), hex_color)
        shd.set(qn("w:val"),  "clear")
        tcPr.append(shd)

    def add_heading(text: str, level: int, color_hex: str = "1a56db"):
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
            run.font.color.rgb = RGBColor(r, g, b)
        return p

    def add_table(headers: list, rows: list, col_widths=None):
        t = doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = "Table Grid"
        # Header row
        hrow = t.rows[0]
        for i, h in enumerate(headers):
            cell = hrow.cells[i]
            cell.text = h
            set_cell_bg(cell, "1a56db")
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(9)
        # Data rows
        for ri, row in enumerate(rows):
            drow = t.rows[ri + 1]
            bg   = "f8f9fa" if ri % 2 == 0 else "ffffff"
            for ci, val in enumerate(row):
                cell = drow.cells[ci]
                cell.text = str(val)
                set_cell_bg(cell, bg)
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(8)
        return t

    # ── Title ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_para.add_run("MALYZE — Malware Analysis Report")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(26, 86, 219)

    # Threat level
    tl_para = doc.add_paragraph()
    run = tl_para.add_run(f"THREAT LEVEL: {threat_level}")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(*threat_rgb)

    doc.add_paragraph()

    # Meta info
    add_heading("Report Metadata", 1)
    add_table(
        ["Property", "Value"],
        [
            ["Analyst",      meta.get("analyst", "N/A")],
            ["Organization", meta.get("org", "N/A")],
            ["Date",         meta.get("timestamp", "N/A")],
            ["Tool",         meta.get("tool", "Malyze")],
            ["Generated",    datetime.datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC")],
        ],
    )
    doc.add_paragraph()

    # File info
    add_heading("File Information", 1)
    add_table(
        ["Property", "Value"],
        [
            ["Name",       file_info.get("name", "N/A")],
            ["Type",       file_info.get("type", "N/A")],
            ["Extension",  file_info.get("extension", "N/A")],
            ["Size",       f"{hashes.get('size',0):,} bytes"],
            ["MD5",        hashes.get("md5", "N/A")],
            ["SHA1",       hashes.get("sha1", "N/A")],
            ["SHA256",     hashes.get("sha256", "N/A")],
            ["Magic Bytes",file_info.get("magic_bytes", "N/A")],
        ],
    )
    doc.add_paragraph()

    # AI Analysis
    add_heading("AI-Powered Threat Analysis", 1)
    if ai.get("error"):
        p = doc.add_paragraph(f"Error: {ai['error']}")
        p.runs[0].font.color.rgb = RGBColor(220, 38, 38)
    elif ai.get("analysis"):
        for line in ai["analysis"].splitlines():
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
            if line.startswith("##"):
                add_heading(line.lstrip("#").strip(), 2)
            elif line.startswith("#"):
                add_heading(line.lstrip("#").strip(), 1)
            else:
                p = doc.add_paragraph(line)
                p.runs[0].font.size = Pt(9) if p.runs else None
    doc.add_paragraph()

    # Entropy
    entropy = static.get("entropy", {})
    if entropy:
        add_heading("Entropy Analysis", 1)
        add_table(
            ["Metric", "Value"],
            [
                ["Overall Entropy",      str(entropy.get("overall_entropy", "N/A"))],
                ["Classification",       entropy.get("classification", "N/A")],
                ["Suspicious",           str(entropy.get("suspicious", False))],
                ["High Entropy Blocks",  f"{entropy.get('high_entropy_blocks',0)} / {entropy.get('total_blocks',0)}"],
            ],
        )
        doc.add_paragraph()

    # Packer
    packer = static.get("packer", {})
    if packer:
        add_heading("Packer / Protector Detection", 1)
        pkrs = packer.get("detected_packers", [])
        p = doc.add_paragraph()
        run = p.add_run("Detected: " + (", ".join(pkrs) if pkrs else "None"))
        run.font.size = Pt(9)
        if pkrs:
            run.font.color.rgb = RGBColor(220, 38, 38)
            run.bold = True
        doc.add_paragraph()

    # PE
    pe = static.get("pe", {})
    if pe and not pe.get("error"):
        add_heading("PE File Analysis", 1)
        add_table(
            ["Field", "Value"],
            [
                ["Machine",       pe.get("machine", "N/A")],
                ["Compile Time",  pe.get("timestamp", "N/A")],
                ["Subsystem",     pe.get("subsystem", "N/A")],
                ["Entry Point",   pe.get("entry_point", "N/A")],
                ["Image Base",    pe.get("image_base", "N/A")],
                ["Is DLL",        str(pe.get("is_dll", False))],
                ["Is 64-bit",     str(pe.get("is_64bit", False))],
                ["Has TLS",       str(pe.get("has_tls", False))],
                ["Has Signature", str(pe.get("has_signature", False))],
            ],
        )

        # Sections
        add_heading("Sections", 2)
        sec_rows = [
            [s.get("name","?"), s.get("virtual_address",""), f"{s.get('virtual_size',0):,}",
             f"{s.get('raw_size',0):,}", f"{s.get('entropy',0):.3f}", ", ".join(s.get("flags",[])[:3])]
            for s in pe.get("sections", [])
        ]
        add_table(["Name","Virt. Addr","V.Size","Raw Size","Entropy","Flags"], sec_rows)

        susp = pe.get("suspicious_imports", [])
        if susp:
            add_heading(f"Suspicious Imports ({len(susp)})", 2)
            p = doc.add_paragraph(", ".join(susp[:60]))
            p.runs[0].font.color.rgb = RGBColor(220, 38, 38)
            p.runs[0].font.size = Pt(8)

        all_imp = pe.get("imports", {})
        if all_imp:
            add_heading("Import Table", 2)
            imp_rows = [
                [dll, str(len(funcs)), ", ".join(funcs[:5]) + ("…" if len(funcs)>5 else "")]
                for dll, funcs in list(all_imp.items())[:20]
            ]
            add_table(["DLL", "Count", "Functions (sample)"], imp_rows)

        exp = pe.get("exports", [])
        if exp:
            add_heading("Exports", 2)
            add_table(
                ["Name", "Ordinal", "Address"],
                [[e.get("name","?"), str(e.get("ordinal","")), e.get("address","")] for e in exp[:30]],
            )
        doc.add_paragraph()

    # IOCs
    strings_data = static.get("strings", {})
    iocs = strings_data.get("iocs", {})
    if iocs:
        add_heading(f"Strings & IOCs ({strings_data.get('total',0)} total)", 1)
        for cat, items in iocs.items():
            add_heading(cat.replace("_", " ").upper(), 2)
            p = doc.add_paragraph(", ".join(str(i) for i in items[:30]))
            if p.runs:
                p.runs[0].font.size = Pt(8)
        doc.add_paragraph()

    # Disassembly
    disasm = static.get("disassembly", {})
    if disasm and not disasm.get("error"):
        add_heading(f"Disassembly ({disasm.get('arch','')})", 1)
        for bname, block in (disasm.get("blocks") or {}).items():
            if block.get("error"):
                continue
            add_heading(f"{bname} @ {block.get('va','')}", 2)
            asm_lines = [
                f"{i['address']}  {i['mnemonic']:<10} {i['op_str']}"
                for i in block.get("instructions", [])[:60]
            ]
            p = doc.add_paragraph("\n".join(asm_lines))
            if p.runs:
                p.runs[0].font.name = "Courier New"
                p.runs[0].font.size = Pt(7)
        doc.add_paragraph()

    # Dynamic
    if dynamic:
        add_heading("Dynamic Analysis", 1)
        exec_r = dynamic.get("execution", {})
        if exec_r:
            add_table(
                ["Field", "Value"],
                [
                    ["PID",       str(exec_r.get("pid", "N/A"))],
                    ["Exit Code", str(exec_r.get("exit_code", "N/A"))],
                    ["Timed Out", str(exec_r.get("timed_out", False))],
                ],
            )

    doc.save(path)
    return path
